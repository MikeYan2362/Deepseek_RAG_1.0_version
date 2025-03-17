import os
import uuid
import time
import chromadb
from chromadb.config import Settings
import re
from typing import List, Dict, Any, Tuple, Optional
import numpy as np
import logging
import hashlib

logger = logging.getLogger(__name__)

# 自定义文本分割器，替代langchain.text_splitter.RecursiveCharacterTextSplitter
class CustomTextSplitter:
    def __init__(self, chunk_size=500, chunk_overlap=50, separators=None):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = separators or ["\n\n", "\n", "。", "！", "？", "：", "；", "，", " ", ""]
    
    def split_text(self, text):
        """将文本分割成块"""
        if not text:
            return []
        
        # 使用分隔符递归地分割文本
        chunks = self._split_by_separators(text, 0)
        
        # 将较小的块合并成大致chunk_size大小的块
        result = []
        current_chunk = []
        current_chunk_size = 0
        
        for chunk in chunks:
            if len(chunk) > self.chunk_size:
                # 如果单个块太大，按字符分割
                if current_chunk:
                    result.append("".join(current_chunk))
                    current_chunk = []
                    current_chunk_size = 0
                
                for i in range(0, len(chunk), self.chunk_size - self.chunk_overlap):
                    result.append(chunk[i:i + self.chunk_size])
            else:
                if current_chunk_size + len(chunk) > self.chunk_size:
                    result.append("".join(current_chunk))
                    # 保留重叠部分
                    overlap_size = min(self.chunk_overlap, current_chunk_size)
                    current_chunk = current_chunk[-overlap_size:] if overlap_size > 0 else []
                    current_chunk_size = sum(len(c) for c in current_chunk)
                
                current_chunk.append(chunk)
                current_chunk_size += len(chunk)
        
        if current_chunk:
            result.append("".join(current_chunk))
        
        return result
    
    def _split_by_separators(self, text, separator_idx):
        """使用指定的分隔符递归分割文本"""
        if separator_idx >= len(self.separators):
            return [text]
        
        separator = self.separators[separator_idx]
        if not separator:
            return [text]
        
        # 使用当前分隔符分割
        splits = text.split(separator)
        
        # 递归地使用下一个分隔符分割
        result = []
        for split in splits:
            if split:
                subsplits = self._split_by_separators(split, separator_idx + 1)
                result.extend(subsplits)
        
        return result
    
    def split_documents(self, documents):
        """分割文档并保留元数据"""
        # 简单实现，假设documents是包含文本和元数据的对象列表
        results = []
        for doc in documents:
            # 假设每个文档有page_content和metadata属性
            if hasattr(doc, 'page_content') and hasattr(doc, 'metadata'):
                for chunk in self.split_text(doc.page_content):
                    results.append(Document(page_content=chunk, metadata=doc.metadata.copy()))
            else:
                # 处理简单字符串
                for chunk in self.split_text(doc):
                    results.append(Document(page_content=chunk, metadata={}))
        return results

# 简单文档类，模拟langchain中的Document
class Document:
    def __init__(self, page_content, metadata=None):
        self.page_content = page_content
        self.metadata = metadata or {}

class DocumentProcessor:
    def __init__(self, persist_directory: str = "./chroma_db"):
        self.persist_directory = persist_directory
        os.makedirs(persist_directory, exist_ok=True)
        
        self.client = chromadb.PersistentClient(
            path=persist_directory,
            settings=Settings(anonymized_telemetry=False)
        )
        
        # 文本分割器
        self.text_splitter = CustomTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            separators=["\n\n", "\n", "。", "！", "？", "：", "；", "，", " ", ""]
        )
    
    def _get_loader_for_file(self, file_path: str):
        """根据文件类型选择适当的加载器"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return self._load_pdf(file_path)
        elif file_extension == '.md':
            return self._load_markdown(file_path)
        elif file_extension in ['.docx', '.doc']:
            return self._load_docx(file_path)
        elif file_extension in ['.txt', '.csv', '.log']:
            return self._load_text(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {file_extension}")
    
    def _load_text(self, file_path):
        """加载文本文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return [Document(page_content=text, metadata={"source": file_path})]
    
    def _load_pdf(self, file_path):
        """加载PDF文件"""
        import fitz  # PyMuPDF
        
        doc = fitz.open(file_path)
        documents = []
        
        for i, page in enumerate(doc):
            text = page.get_text()
            if text.strip():  # 仅添加非空页面
                documents.append(Document(
                    page_content=text,
                    metadata={"source": file_path, "page": i + 1}
                ))
        
        return documents
    
    def _load_markdown(self, file_path):
        """加载Markdown文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return [Document(page_content=text, metadata={"source": file_path})]
    
    def _load_docx(self, file_path):
        """加载DOCX文件"""
        try:
            import docx
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return [Document(page_content=text, metadata={"source": file_path})]
        except ImportError:
            logger.warning("python-docx未安装，无法处理DOCX文件")
            return [Document(page_content="", metadata={"source": file_path, "error": "docx加载失败"})]
    
    def _get_or_create_collection(self, collection_name: str):
        """获取或创建一个集合"""
        try:
            return self.client.get_collection(name=collection_name)
        except:
            return self.client.create_collection(name=collection_name)
    
    def _simple_text_to_vector(self, text: str) -> List[float]:
        """简单文本向量化：使用MD5哈希并规范化数值
           这不是真正语义的向量化，但在不使用复杂模型的情况下可以用于测试"""
        # 使用MD5哈希，确保相同的文本生成相同的向量
        text = text.lower()  # 简单的标准化
        hash_obj = hashlib.md5(text.encode())
        # 将哈希值转换为1536维向量（大致接近OpenAI的embedding大小）
        hex_digest = hash_obj.hexdigest()
        
        # 生成一个足够长的向量 (1536 维)
        vec_size = 1536
        step = len(hex_digest) // vec_size
        if step == 0:
            # 如果hex_digest不够长，重复它
            hex_digest = hex_digest * (vec_size // len(hex_digest) + 1)
            step = len(hex_digest) // vec_size
            
        # 创建向量
        vector = []
        for i in range(0, len(hex_digest) - step, step):
            # 将十六进制字符转换为浮点数
            hex_value = hex_digest[i:i+step]
            # 归一化到 [-1, 1] 范围
            float_value = (int(hex_value, 16) / (16 ** len(hex_value))) * 2 - 1
            vector.append(float_value)
        
        # 确保向量大小正确
        vector = vector[:vec_size]
        while len(vector) < vec_size:
            vector.append(0.0)
            
        return vector
    
    def _generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """使用简单的文本哈希生成向量嵌入"""
        return [self._simple_text_to_vector(text) for text in texts]
    
    async def process_file(self, file_path: str, knowledge_base_id: str) -> Dict[str, Any]:
        """处理文件并添加到向量数据库"""
        try:
            start_time = time.time()
            
            # 1. 加载文档
            documents = self._get_loader_for_file(file_path)
            
            # 2. 分割文档
            chunks = self.text_splitter.split_documents(documents)
            
            # 3. 获取集合
            collection = self._get_or_create_collection(knowledge_base_id)
            
            # 4. 处理每个文本块
            document_id = str(uuid.uuid4())
            file_name = os.path.basename(file_path)
            
            ids = []
            texts = []
            metadatas = []
            
            for i, chunk in enumerate(chunks):
                chunk_id = f"{document_id}_chunk_{i}"
                ids.append(chunk_id)
                texts.append(chunk.page_content)
                
                # 元数据
                metadata = {
                    "document_id": document_id,
                    "document_name": file_name,
                    "chunk_index": i
                }
                
                # 添加页码信息（如果有）
                if 'page' in chunk.metadata:
                    metadata['page'] = chunk.metadata['page']
                
                metadatas.append(metadata)
            
            # 5. 生成嵌入并添加到数据库
            embeddings = self._generate_embeddings(texts)
            
            collection.add(
                ids=ids,
                embeddings=embeddings,
                documents=texts,
                metadatas=metadatas
            )
            
            process_time = time.time() - start_time
            
            return {
                "document_id": document_id,
                "document_name": file_name,
                "chunks_count": len(chunks),
                "process_time_seconds": process_time
            }
            
        except Exception as e:
            logger.error(f"处理文件失败: {str(e)}")
            raise
    
    async def semantic_search(self, 
                             query: str, 
                             knowledge_base_id: str,
                             top_k: int = 3) -> List[Dict[str, Any]]:
        """语义搜索"""
        try:
            # 1. 获取集合
            collection = self._get_or_create_collection(knowledge_base_id)
            
            # 2. 生成查询嵌入
            query_embedding = self._simple_text_to_vector(query)
            
            # 3. 执行搜索
            results = collection.query(
                query_embeddings=[query_embedding],
                n_results=top_k,
                include=["documents", "metadatas", "distances"]
            )
            
            # 4. 格式化结果
            formatted_results = []
            if not results['documents']:
                return formatted_results
                
            for i, (doc, metadata, distance) in enumerate(zip(
                results['documents'][0],
                results['metadatas'][0],
                results['distances'][0]
            )):
                formatted_results.append({
                    "content": doc,
                    "document_id": metadata.get("document_id", ""),
                    "document_name": metadata.get("document_name", ""),
                    "page": metadata.get("page"),
                    "similarity": 1.0 - (distance / 2.0)  # 将距离转换为相似度
                })
            
            return formatted_results
            
        except Exception as e:
            logger.error(f"语义搜索失败: {str(e)}")
            raise 